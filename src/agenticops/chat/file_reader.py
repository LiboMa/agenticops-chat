"""File content readers for chat context injection."""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Supported extensions
TEXT_EXTENSIONS = {
    ".txt", ".log", ".md", ".json", ".yaml", ".yml", ".toml",
    ".ini", ".cfg", ".conf", ".sh", ".py", ".js", ".ts", ".csv",
    ".xml", ".html", ".tf", ".hcl",
}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}

MAX_FILE_SIZE = 512 * 1024  # 512 KB text limit
MAX_IMAGE_SIZE = 5 * 1024 * 1024  # 5 MB image limit
MAX_DOCUMENT_SIZE = 5 * 1024 * 1024  # 5 MB document limit

# Strands SDK native format maps (for multimodal content blocks)
IMAGE_FORMAT_MAP: dict[str, str] = {
    ".png": "png", ".jpg": "jpeg", ".jpeg": "jpeg",
    ".gif": "gif", ".webp": "webp",
}

DOCUMENT_FORMAT_MAP: dict[str, str] = {
    ".pdf": "pdf", ".csv": "csv", ".doc": "doc", ".docx": "docx",
    ".xls": "xls", ".xlsx": "xlsx", ".html": "html", ".txt": "txt", ".md": "md",
}


def read_file_as_text(path: str) -> tuple[str, str | None]:
    """Read a file and return (content_string, error_or_none).

    For text files: returns raw content.
    For DOCX: extracts text via python-docx.
    For PDF: extracts text via pymupdf or pypdf.
    For images: returns a placeholder description.
    """
    p = Path(path).expanduser().resolve()

    if not p.exists():
        return "", f"File not found: {path}"

    if not p.is_file():
        return "", f"Not a file: {path}"

    suffix = p.suffix.lower()
    file_size = p.stat().st_size

    # Text files
    if suffix in TEXT_EXTENSIONS or suffix == "":
        if file_size > MAX_FILE_SIZE:
            return "", f"File too large ({file_size} bytes, max {MAX_FILE_SIZE}): {path}"
        try:
            content = p.read_text(encoding="utf-8", errors="replace")
            return content, None
        except Exception as e:
            return "", f"Error reading {path}: {e}"

    # DOCX
    if suffix in DOCX_EXTENSIONS:
        try:
            from docx import Document
        except ImportError:
            return "", "python-docx not installed. Run: pip install python-docx"
        try:
            doc = Document(str(p))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs), None
        except Exception as e:
            return "", f"Error reading DOCX {path}: {e}"

    # PDF
    if suffix in PDF_EXTENSIONS:
        try:
            import pymupdf

            doc = pymupdf.open(str(p))
            pages = [page.get_text() for page in doc]
            return "\n\n".join(pages), None
        except ImportError:
            pass
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(p))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages), None
        except ImportError:
            return "", "No PDF library installed. Run: pip install pymupdf"
        except Exception as e:
            return "", f"Error reading PDF {path}: {e}"

    # Images — describe presence
    if suffix in IMAGE_EXTENSIONS:
        if file_size > MAX_IMAGE_SIZE:
            return "", f"Image too large ({file_size} bytes, max {MAX_IMAGE_SIZE}): {path}"
        return (
            f"[Image file: {p.name}, size: {file_size} bytes, type: {suffix}. "
            f"Image content cannot be directly analyzed in text mode. "
            f"Please describe what you need analyzed about this image.]"
        ), None

    return "", f"Unsupported file type: {suffix}"


def is_image_file(path_or_name: str) -> bool:
    """Check if the file extension is a Strands SDK-supported image format."""
    return Path(path_or_name).suffix.lower() in IMAGE_FORMAT_MAP


def is_document_file(path_or_name: str) -> bool:
    """Check if the file extension is a Strands SDK-supported document format."""
    return Path(path_or_name).suffix.lower() in DOCUMENT_FORMAT_MAP


def read_file_as_image_bytes(path: str) -> tuple[bytes | None, str | None, str | None]:
    """Read an image file and return raw bytes for Strands ContentBlock.

    Returns (raw_bytes, format, error).
    """
    p = Path(path).expanduser().resolve()
    if not p.exists():
        return None, None, f"File not found: {path}"
    if not p.is_file():
        return None, None, f"Not a file: {path}"

    suffix = p.suffix.lower()
    fmt = IMAGE_FORMAT_MAP.get(suffix)
    if not fmt:
        return None, None, f"Unsupported image format: {suffix}"

    file_size = p.stat().st_size
    if file_size > MAX_IMAGE_SIZE:
        return None, None, f"Image too large ({file_size} bytes, max {MAX_IMAGE_SIZE}): {path}"

    try:
        return p.read_bytes(), fmt, None
    except Exception as e:
        return None, None, f"Error reading image {path}: {e}"


def read_upload_image_bytes(
    filename: str, raw_bytes: bytes,
) -> tuple[bytes | None, str | None, str | None]:
    """Process uploaded image bytes for Strands ContentBlock.

    Returns (raw_bytes, format, error).
    """
    suffix = Path(filename).suffix.lower()
    fmt = IMAGE_FORMAT_MAP.get(suffix)
    if not fmt:
        return None, None, f"Unsupported image format: {suffix}"
    if len(raw_bytes) > MAX_IMAGE_SIZE:
        return None, None, f"Image too large ({len(raw_bytes)} bytes, max {MAX_IMAGE_SIZE})"
    return raw_bytes, fmt, None


def read_file_as_document_bytes(
    path: str,
) -> tuple[bytes | None, str | None, str | None, str | None]:
    """Read a document file and return raw bytes for Strands ContentBlock.

    Returns (raw_bytes, format, name, error).
    """
    p = Path(path).expanduser().resolve()
    if not p.exists():
        return None, None, None, f"File not found: {path}"
    if not p.is_file():
        return None, None, None, f"Not a file: {path}"

    suffix = p.suffix.lower()
    fmt = DOCUMENT_FORMAT_MAP.get(suffix)
    if not fmt:
        return None, None, None, f"Unsupported document format: {suffix}"

    file_size = p.stat().st_size
    if file_size > MAX_DOCUMENT_SIZE:
        return None, None, None, f"Document too large ({file_size} bytes, max {MAX_DOCUMENT_SIZE}): {path}"

    try:
        return p.read_bytes(), fmt, p.name, None
    except Exception as e:
        return None, None, None, f"Error reading document {path}: {e}"


def read_upload_document_bytes(
    filename: str, raw_bytes: bytes,
) -> tuple[bytes | None, str | None, str | None, str | None]:
    """Process uploaded document bytes for Strands ContentBlock.

    Returns (raw_bytes, format, name, error).
    """
    suffix = Path(filename).suffix.lower()
    fmt = DOCUMENT_FORMAT_MAP.get(suffix)
    if not fmt:
        return None, None, None, f"Unsupported document format: {suffix}"
    if len(raw_bytes) > MAX_DOCUMENT_SIZE:
        return None, None, None, f"Document too large ({len(raw_bytes)} bytes, max {MAX_DOCUMENT_SIZE})"
    return raw_bytes, fmt, filename, None


def read_upload_bytes(filename: str, raw_bytes: bytes) -> tuple[str, str | None]:
    """Extract text content from uploaded file bytes (for web uploads).

    Returns (content_string, error_or_none).
    """
    suffix = Path(filename).suffix.lower()

    # Text files
    if suffix in TEXT_EXTENSIONS or suffix == "":
        if len(raw_bytes) > MAX_FILE_SIZE:
            return "", f"File too large ({len(raw_bytes)} bytes, max {MAX_FILE_SIZE})"
        return raw_bytes.decode("utf-8", errors="replace"), None

    # DOCX
    if suffix in DOCX_EXTENSIONS:
        try:
            from docx import Document
            from io import BytesIO

            doc = Document(BytesIO(raw_bytes))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text, None
        except ImportError:
            return "", "DOCX support not installed on server"
        except Exception as e:
            return "", f"Failed to read DOCX: {e}"

    # PDF
    if suffix in PDF_EXTENSIONS:
        try:
            import pymupdf

            doc = pymupdf.open(stream=raw_bytes, filetype="pdf")
            text = "\n\n".join(page.get_text() for page in doc)
            return text, None
        except ImportError:
            pass
        try:
            from pypdf import PdfReader
            from io import BytesIO

            reader = PdfReader(BytesIO(raw_bytes))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            return text, None
        except ImportError:
            return "", "PDF support not installed on server"
        except Exception as e:
            return "", f"Failed to read PDF: {e}"

    # Images
    if suffix in IMAGE_EXTENSIONS:
        return (
            f"[Uploaded image: {filename}, {len(raw_bytes)} bytes. "
            f"Image analysis not supported in text mode.]"
        ), None

    return "", f"Unsupported file type: {suffix}"
